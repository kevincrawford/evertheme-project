import io
import re
from math import ceil
from pathlib import Path

from app.schemas.document import DocumentHints

SUPPORTED_TYPES = {".docx", ".pdf", ".txt", ".md"}

_LARGE_DOC_THRESHOLD = 20_000   # chars (~10 pages / ~5 000 words)
_FALLBACK_CHUNK_SIZE = 12_000   # used only for estimated_chunks in hints


def parse_document(content: bytes, filename: str) -> str:
    """Extract plain text with embedded [§ Section] markers from an uploaded document."""
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_TYPES:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: {', '.join(SUPPORTED_TYPES)}")

    if suffix == ".docx":
        return _parse_docx(content)
    elif suffix == ".pdf":
        return _parse_pdf(content)
    elif suffix == ".md":
        return _parse_markdown(content)
    else:
        return _parse_txt(content)


def analyse_document(text: str, file_type: str) -> DocumentHints:
    """Compute advisory metadata from already-extracted text. No LLM call needed."""
    char_count = len(text)
    section_count = len(re.findall(r"^\[§ .+\]$", text, re.MULTILINE))
    is_large = char_count > _LARGE_DOC_THRESHOLD
    estimated_chunks = max(1, ceil(char_count / _FALLBACK_CHUNK_SIZE))

    format_warnings: list[str] = []
    if file_type == ".txt":
        format_warnings.append(
            "Plain text has limited section detection. "
            "Converting to .docx will improve story grouping accuracy."
        )
    elif file_type == ".pdf":
        format_warnings.append(
            "PDF section detection uses heuristics and may miss some boundaries. "
            ".docx provides the most accurate results."
        )

    processing_note: str | None = None
    if is_large:
        processing_note = (
            f"This document will be processed in ~{estimated_chunks} section(s). "
            "Generation may take longer than usual."
        )

    return DocumentHints(
        char_count=char_count,
        section_count=section_count,
        is_large=is_large,
        estimated_chunks=estimated_chunks,
        format_warnings=format_warnings,
        processing_note=processing_note,
    )


# ── Format-specific parsers ───────────────────────────────────────────────────

def _parse_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        style = p.style.name if p.style else ""
        if style.startswith("Heading"):
            parts.append(f"[§ {p.text.strip()}]")
        else:
            parts.append(p.text)

    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    return "\n\n".join(parts)


def _parse_pdf(content: bytes) -> str:
    import pdfplumber
    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    raw = "\n\n".join(pages)
    return _inject_heuristic_section_markers(raw)


def _parse_markdown(content: bytes) -> str:
    text = content.decode("utf-8", errors="replace")
    lines = text.splitlines()
    out: list[str] = []
    for line in lines:
        m = re.match(r"^(#{1,3})\s+(.+)", line)
        if m:
            out.append(f"[§ {m.group(2).strip()}]")
        else:
            out.append(line)
    return "\n".join(out)


def _parse_txt(content: bytes) -> str:
    text = content.decode("utf-8", errors="replace")
    return _inject_heuristic_section_markers(text)


def _inject_heuristic_section_markers(text: str) -> str:
    """
    Heuristically detect section headings in plain text / PDF output and
    replace them with [§ Heading] markers.

    Detects:
    - Numbered sections:  "1. Introduction" / "2.1 Functional Requirements"
    - ALL-CAPS short lines (≤ 60 chars) standing alone
    - Short lines (< 80 chars) followed by a blank line and not ending with "."
    """
    lines = text.splitlines()
    out: list[str] = []

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            out.append(line)
            continue

        is_numbered = bool(re.match(r"^\d+(\.\d+)*\.?\s+[A-Z]", stripped))
        is_all_caps = (
            stripped == stripped.upper()
            and len(stripped) <= 60
            and len(stripped) > 3
            and not stripped[-1].isdigit()
        )
        next_blank = i + 1 < len(lines) and not lines[i + 1].strip()
        is_short_heading = (
            len(stripped) < 80
            and next_blank
            and not stripped.endswith(".")
            and not stripped.endswith(",")
        )

        if is_numbered or is_all_caps or is_short_heading:
            heading = re.sub(r"[.:]\s*$", "", stripped)
            out.append(f"[§ {heading}]")
        else:
            out.append(line)

    return "\n".join(out)
